#from dateutil.relativedelta import relativedelta
#from django.db.models import Max
import csv
import io
import subprocess

from django.shortcuts import render, get_object_or_404, redirect
#from django.utils import timezone
from datetime import datetime

from docx2pdf import convert

from .models import Equipment, MaintenanceRecord
from .filters import EquipmentFilter
from .forms import EquipmentForm, MaintenanceRecordForm
from .services import MaintenanceService, ReportService, CalendarService  # Включаем сервисы
import os
from django.http import HttpResponse
from docx import Document
from django.conf import settings
import pypandoc

def dashboard(request):
    """Отображение дашборда с просроченными и предстоящими работами."""
    days = int(request.GET.get('days', 30))  # Количество дней по умолчанию — 30
    upcoming_maintenance, overdue_maintenance = MaintenanceService.get_upcoming_and_overdue_maintenance(days)

    equipment_filter = EquipmentFilter(request.GET, queryset=Equipment.objects.all())
    return render(request, 'equipment/dashboard.html', {
        'upcoming_maintenance': upcoming_maintenance,
        'overdue_maintenance': overdue_maintenance,
        'days': days,
        'filter': equipment_filter,
    })


def equipment_list(request):
    """Отображение таблицы оборудования с фильтрацией и сортировкой."""
    equipment_filter = EquipmentFilter(request.GET, queryset=Equipment.objects.all())
    return render(request, 'equipment/equipment_list.html', {'filter': equipment_filter})


def equipment_detail(request, pk):
    """Отображение детальной информации о конкретном оборудовании."""
    equipment = get_object_or_404(Equipment, pk=pk)
    maintenance_records = MaintenanceRecord.objects.filter(equipment=equipment).select_related('maintenance_type')
    return render(request, 'equipment/equipment_detail.html', {
        'equipment': equipment,
        'maintenance_records': maintenance_records,
    })


def equipment_create(request):
    """Создание нового оборудования."""
    if request.method == 'POST':
        form = EquipmentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('dashboard')
    else:
        form = EquipmentForm()
    return render(request, 'equipment/equipment_form.html', {'form': form})


def equipment_edit(request, pk):
    """Редактирование существующего оборудования."""
    equipment = get_object_or_404(Equipment, pk=pk)
    if request.method == 'POST':
        form = EquipmentForm(request.POST, instance=equipment)
        if form.is_valid():
            form.save()
            return redirect('equipment_detail', pk=pk)
    else:
        form = EquipmentForm(instance=equipment)
    return render(request, 'equipment/equipment_form.html', {'form': form})


def maintenance_create(request, equipment_id):
    """Добавление записи обслуживания для оборудования."""
    equipment = get_object_or_404(Equipment, id=equipment_id)
    if request.method == 'POST':
        form = MaintenanceRecordForm(request.POST, request.FILES)
        if form.is_valid():
            maintenance_record = MaintenanceService.create_maintenance(form, equipment)
            return redirect('equipment_detail', pk=equipment.id)
    else:
        form = MaintenanceRecordForm()
    return render(request, 'equipment/add_maintenance.html', {'form': form, 'equipment': equipment})


def maintenance_edit(request, pk):
    """Редактирование записи обслуживания."""
    maintenance = get_object_or_404(MaintenanceRecord, pk=pk)
    if request.method == 'POST':
        form = MaintenanceRecordForm(request.POST, request.FILES, instance=maintenance)
        if form.is_valid():
            form.save()
            return redirect('equipment_detail', pk=maintenance.equipment.pk)
    else:
        form = MaintenanceRecordForm(instance=maintenance)
    return render(request, 'equipment/maintenance_form.html', {'form': form, 'equipment': maintenance.equipment})


def generate_equipment_report(request):
    """Генерация отчета по оборудованию на основе даты покупки."""
    date = request.GET.get('date', datetime.now().strftime('%Y-%m-%d'))
    equipments = ReportService.generate_equipment_report(date)
    return render(request, 'equipment/equipment_report.html', {
        'equipments': equipments,
        'report_date': date,
    })


def generate_maintenance_calendar(request):
    """Генерация календарного графика обслуживания по каждому виду обслуживания."""
    year = int(request.GET.get('year', datetime.now().year))
    maintenance_calendars = CalendarService.generate_maintenance_calendar(year)
    return render(request, 'equipment/maintenance_calendar.html', {
        'maintenance_calendars': maintenance_calendars,
        'year': year,
        'month_numbers': range(1, 13),
    })

def replace_placeholder(paragraph, equipment):
    """Заменяет плейсхолдеры в тексте параграфа."""
    if '{{ name }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ name }}', equipment.name)
    if '{{ model }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ model }}', equipment.model)
    if '{{ serial_number }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ serial_number }}', equipment.serial_number)
    if '{{ code }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ code }}', equipment.code)
    if '{{ manufacturer }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ manufacturer }}', equipment.manufacturer)
    if '{{ purchase_date }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ purchase_date }}', str(equipment.purchase_date))
    if '{{ start_date }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ start_date }}', str(equipment.start_date))
    if '{{ responsible_person }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ responsible_person }}', equipment.responsible_person)
    if '{{ condition }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ condition }}', equipment.condition)
    if '{{ service_contact }}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ service_contact }}', equipment.service_contact)

def generate_equipment_pdf(request, equipment_id):
    equipment = Equipment.objects.get(id=equipment_id)

    # Путь к шаблону .docx
    template_path = os.path.join(settings.BASE_DIR, 'templates/docx/equipment_template.docx')
    document = Document(template_path)

    # Заменяем плейсхолдеры в параграфах
    for paragraph in document.paragraphs:
        replace_placeholder(paragraph, equipment)

    # Обрабатываем таблицы
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder(paragraph, equipment)

    # Сохраняем заполненный .docx временно
    docx_path = os.path.join(settings.MEDIA_ROOT, f'equipment_{equipment_id}.docx')
    document.save(docx_path)

    # Путь для сохранения PDF
    pdf_path = os.path.join(settings.MEDIA_ROOT, f'equipment_{equipment_id}.pdf')

    # Укажите полный путь к LibreOffice (замените на правильный путь к LibreOffice)
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

    # Используем LibreOffice для конвертации .docx в PDF
    subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', settings.MEDIA_ROOT])

    # Возврат PDF файла как ответа для отображения в браузере
    with open(pdf_path, 'rb') as pdf_file:
        response = HttpResponse(pdf_file.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'inline; filename="equipment_{equipment_id}.pdf"'

    # Удаление временных файлов
    os.remove(docx_path)
    os.remove(pdf_path)

    return response

def import_csv(request):
    if request.method == 'POST':
        csv_file = request.FILES['csv_file']

        # Проверка на тип файла (если нужно)
        if not csv_file.name.endswith('.csv'):
            return render(request, 'import.html', {'error': 'Неправильный формат файла. Пожалуйста, загрузите CSV файл.'})

        # Чтение и обработка CSV-файла
        data_set = csv_file.read().decode('UTF-8')
        io_string = io.StringIO(data_set)
        csv_reader = csv.reader(io_string, delimiter=';')

        # Пропускаем заголовок
        next(csv_reader)

        for row in csv_reader:
            _, created = Equipment.objects.update_or_create(
                name=row[0],
                model=row[1],
                manufacturer=row[2],
                serial_number=row[3],
                code=row[4],
                location=row[5],
                responsible_person=row[6],
                purchase_date='2022-12-30',
                start_date='2023-07-31',
                department='Bal',
                condition='Nowe',
                service_contact=' ',
            )

        return render(request, 'import.html', {'success': 'Данные успешно импортированы!'})

    return render(request, 'import.html')
